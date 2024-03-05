import os
import time
from reading import read_excel_vehicular
import numpy as np
import shutil
from openpyxl import load_workbook
from openpyxl.styles import NamedStyle, Border, Side
import re
import docx
from  docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math

def _put_format(table,row,col): #Just bold
    paragraph = table.cell(row,col).paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True
def _join_cells(table,row1,col1,row2,col2):
    table.rows[row1].cells[col1].merge(table.rows[row2].cells[col2])
def vehicular(entregable_path):

    hojas = ["N","S","E","O"]

    ruta = entregable_path
    listas = [[] for _ in range(4)]

    ruta_vehicular = os.path.join(ruta,"7.- Informacion de Campo","Vehicular")
    files = ["Atipico",
            "Atipico (Protransito)",
            "Tipico",
            "Tipico (Protransito)",]

    for file in files:
        ruta_tipicidad = os.path.join(ruta_vehicular,file)
        list_excels = os.listdir(ruta_tipicidad)

        for excel in list_excels:
            path_excel = os.path.join(ruta_tipicidad,excel)

            if not "~$" in excel:
                if file == "Tipico":
                    listas[0].append(path_excel)
                elif file == "Atipico":
                    listas[1].append(path_excel)
                elif file == "Tipico (Protransito)":
                    listas[2].append(path_excel)
                elif file == "Atipico (Protransito)":
                    listas[3].append(path_excel)

    archivos = [[] for _ in range(2)]

    #Comparison among Tipico
    for veh2 in listas[2]: #Supervisor
        for veh0 in listas[0]: #Consultor
            nombre_veh2 = os.path.basename(veh2)
            nombre_veh0 = os.path.basename(veh0)

            if nombre_veh2 != nombre_veh0 and nombre_veh2.replace("_REV","") == nombre_veh0:
                archivos[0].append((veh2,veh0))

    #Comparison among Atipico
    for veh3 in listas[3]: #Supervisor
        for veh1 in listas[1]: #Consultor
            nombre_veh3 = os.path.basename(veh3)
            nombre_veh1 = os.path.basename(veh1)

            if nombre_veh3 != nombre_veh1 and nombre_veh3.replace("_REV","") == nombre_veh1:
                archivos[1].append((veh3,veh1))

    data_row = [[] for _ in range(2)]

    inicio_lectura = time.perf_counter()

    ###################
    # Report Creation #
    ###################

    for index, archivo in enumerate(archivos):
        if index == 0:
            print("#### Reportes Tipicos ####")
        else:
            print("#### Reportes Atipicos ####")

        for l,pair in enumerate(archivo):
            if index == 0:
                print(f"Comparando Tipicos ({l+1}/{len(archivo)})")
            else:
                print(f"Comparando Atipicos ({l+1}/{len(archivo)})")

            data_consultor,fecha,giros,_ = read_excel_vehicular(pair[1],False)
            data_supervisor,_,giros_compare,quarter_hour = read_excel_vehicular(pair[0],True)
            
            giros_cons = []
            for giro in giros:
                AUX = [elem for elem in giro if isinstance(elem,int)]
                giros_cons.append(AUX)

            giros_sup = []
            for giro in giros_compare:
                AUX = [elem for elem in giro if isinstance(elem,int)]
                giros_sup.append(AUX)

            stop = 0
            for (i,giro_cons), giro_sup in zip(enumerate(giros_cons), giros_sup):
                if giro_cons == giro_sup:
                    pass
                else:
                    stop += 1
            check = True #Read
            if stop != 0:
                check = False #Not read

            ruta_destino = pair[0].replace("_REV.xlsm","_REP.xlsx")

            container_file = os.path.dirname(pair[1])
            _,tipicidad = os.path.split(container_file)

            resultados = []

            # Reading Info #
            A = []

            for data_con, data_sup in zip(data_consultor, data_supervisor):
                mascara = data_sup != 0
                resultado = np.zeros_like(data_con)
                resultado[mascara] = (data_con[mascara]-data_sup[mascara])
                resultados.append(resultado)

                #Variable para la creación de la tabla de frecuencias
                A.extend(np.abs(resultado[mascara]).tolist())

            path_formato = "./images/Formato_Vehicular.xlsx"
            current_file, name_excel = os.path.split(ruta_destino)
            destiny_file = os.path.join(current_file,'Reportes')
            final_route = os.path.join(destiny_file,name_excel)
            if not os.path.exists(destiny_file):
                os.makedirs(destiny_file)
            shutil.copyfile(path_formato,final_route)
            wb = load_workbook(final_route)

            # Formats #
            porcentaje_style = NamedStyle(name="porcentaje")
            porcentaje_style.border = Border(left=Side(style='thin'),
                                            right=Side(style='thin'),
                                            top=Side(style='thin'),
                                            bottom=Side(style='thin'),)

            for hoja, giro in zip(hojas,giros):
                ws = wb[hoja]
                for p, valor in enumerate(giro):
                    celda = ws.cell(row=15, column=p+11)
                    celda.value = valor

            for hoja, result in zip(hojas,resultados):
                ws = wb[hoja]
                for k, fila in enumerate(result):
                    for l, valor in enumerate(fila):
                        celda = ws.cell(row=k+16, column=l+11)

                        if valor == 0: celda.value = None
                        else: celda.value = valor

                        celda.style = porcentaje_style

            wb.save(final_route)

            _, nombre_archivo = os.path.split(ruta_destino)
            patron1 = r"([A-Z]+[0-9]+)"
            patron2 = r"([A-Z]+-[0-9]+)"
            coincidencia1 = re.search(patron1, nombre_archivo)
            coincidencia2 = re.search(patron2, nombre_archivo)

            if coincidencia1:
                codigo = coincidencia1.group(1)
                codigo = codigo[:2]+'-'+codigo[2:]
            elif coincidencia2:
                codigo = coincidencia2.group(1)
            else:
                print(f"El archivo de excel no posee Código de intersección:\n{nombre_archivo}")
                codigo = "ERROR"

            data_row[index].append((codigo,fecha[:-9],tipicidad,A,quarter_hour,check))
            
    fin_lectura = time.perf_counter()
    duracion_lectura = fin_lectura - inicio_lectura
    print(f"Processing Time in Reading: {duracion_lectura:.2f} seconds.")

    #########################
    # Creating Word Report #
    #########################

    inicio_word = time.perf_counter()

    doc = docx.Document()
    doc.add_heading(f"REPORTE FLUJO VEHICULAR")
    titulos = ['Clase','Frecuencia','Frecuencia Relativa','Frecuencia Acumulada']

    for i, data in enumerate(data_row):
        if len(data) == 0:
            continue

        for dt in data:
            if dt[5]==False:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{dt[0]}: Los giros de Protránsito no coinciden con los giros de EDSA.")
                continue
            
            CODIGO = dt[0]
            FECHA = dt[1]
            TIPICIDAD = dt[2]
            LISTA = dt[3]
            HORA = dt[4]

            #Tratamiento de la lista completa
            num_datos = len(LISTA)
            num_clases = int(1+3.322*(num_datos**(1/3)))
            rango = max(LISTA) - min(LISTA)
            tamaño_clase = rango / num_clases
            tamaño_clase_entero = math.ceil(tamaño_clase)

            #limite_inferior = min(LISTA)
            limite_inferior = 1
            clases = []
            for _ in range(num_clases):
                limite_superior = limite_inferior + tamaño_clase_entero
                clases.append((limite_inferior,limite_superior))
                limite_inferior = limite_superior

            frecuencias = [0]*num_clases
            for dato in LISTA:
                for j, clase in enumerate(clases):
                    if clase[0] <= dato < clase[1]:
                        frecuencias[j] += 1
                        break

            frec_relativa = [f / sum(frecuencias) for f in frecuencias]
            frec_acumulada = [sum(frec_relativa[:i+1]) for i in range(num_clases)]

            table = doc.add_table(rows=1,cols=4,style="Table Grid")
            table.cell(0,0).text = 'Código'
            table.cell(0,1).text = 'Fecha'
            table.cell(0,2).text = 'Tipicidad'
            table.cell(0,3).text = 'Hora'

            for j in range(4):
                _put_format(table,0,j)

            #_join_cells(table,0,2,0,3)

            row = table.add_row().cells
            row[0].text = CODIGO
            row[1].text = str(FECHA)
            row[2].text = str(TIPICIDAD)
            horas_inicio = HORA//4
            minutos_inicio = int(((HORA/4-HORA//4))*60)
            horas_fin = (HORA+1)//4
            minutos_fin = int((((HORA+1)/4-(HORA+1)//4))*60)
            row[3].text = f"{horas_inicio:02}:{minutos_inicio:02} - {horas_fin:02}:{minutos_fin:02}"
            #_join_cells(table,1,2,1,3)

            row = table.add_row().cells
            for i, titulo in enumerate(titulos):
                row[i].text = titulo

            for clase,frecuencia,relativa,acumulada in zip(clases,frecuencias,frec_relativa,frec_acumulada):
                row = table.add_row().cells
                row[0].text = str(f"{int(clase[0])} - {int(clase[1])}")
                row[1].text = str(frecuencia)
                row[2].text = str(int(relativa*100))+'%'
                row[3].text = str(int(acumulada*100))+'%'

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            _put_format(table,2,0)
            _put_format(table,2,1)
            _put_format(table,2,2)
            _put_format(table,2,3)
            
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break()

    try:
        direccion_archivo = os.path.join(ruta_vehicular,f"INFORME_GENERAL_VEHICULAR.docx")
        doc.save(direccion_archivo)
        print("¡Informe general guardado con éxito!")
    except IndexError:
        print(f"No hay datos para generar un informe general")

    final_word = time.perf_counter()
    duracion_word = final_word-inicio_word
    print(f"Processing Time - Word: {duracion_word:.2f} seconds.")
    print("#### Comparación de flujogramas vehiculares finalizado ####")