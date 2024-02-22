from openpyxl import load_workbook
import os
import re
import docx
from docx.shared import Pt,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def times_eyb(ruta):

    ruta_eyb = os.path.join(ruta,"7.- Informacion de Campo","Embarque y Desembarque")
    files = ["Atipico","Tipico"]
    listas = [[] for _ in range(2)]

    for index,file in enumerate(files):
        ruta_tipicidad = os.path.join(ruta_eyb,file)
        try:
            list_excels = os.listdir(ruta_tipicidad)
        except FileNotFoundError:
            print(f"No hay datos en la carpeta de {file}")
            listas[index] = []
            list_excels = []

        for excel in list_excels:
            path_excel = os.path.join(ruta_tipicidad,excel)

            if not "~$" in excel:
                if file == "Tipico":
                    listas[0].append(path_excel)
                else:
                    listas[1].append(path_excel)

    list_counts = []
    list_codigos = []
    for index, lista in enumerate(listas):
        if index == 0:
            print("********** Revisando Tiempos de Emb. y Desemb. Tipicos **********")
        else:
            print("********** Revisando Tiempos de Emb. y Desemb. Atipicos **********")
        
        codigo_excel = []
        count_excel = []
        for index2, excel_path in enumerate(lista):
            print(f"Revisando excel ({index2+1}/{len(lista)})")
            wb = load_workbook(excel_path,read_only=True,data_only=True)
            try:
                ws = wb["Base Data"]
            except KeyError:
                print(f"Error en excel:\n{excel_path}")

            periods_counts = []
            columns = [
                [row[0].value for row in ws[s]]
                for s in [
                    slice("G8","G57"),
                    slice("G58","G107"),
                    slice("G108","G157"),
                ]
            ]

            for column in columns:
                count_none = column.count(None)
                periods_counts.append(count_none) #Un elemento periods_counts contiene los conteos de un excel

            _, name_excel = os.path.split(excel_path)
            #patron = r"([A-Z]+-[0-9]+)"
            patron = r"([A-Z]+[0-9]+)"
            coincidencia = re.search(patron,name_excel)
            if coincidencia:
                codigo = coincidencia.group(1)
                codigo = codigo[:2]+'-'+codigo[2:]
            count_excel.append(periods_counts)
            codigo_excel.append(codigo)
        list_codigos.append(codigo_excel) #Deberia haber solo dos
        list_counts.append(count_excel) #Deberia haber solo dos

    ########################
    # Creating Word Report #
    ########################

    doc = docx.Document()
    doc.add_heading("REPORTE TIEMPOS DE EMBARQUE Y DESEMBARQUE")
    table = doc.add_table(rows=1,cols=6,style="Table Grid")
    section = doc.sections[0]

    for i in range(3): #El 3 se puede cambiar a antojo, pero creo que esta bien.
        for cell in table.columns[i+1].cells:
            cell.width = Inches(1)

    titulos = ['Códigos',
            'Tipicidad',
            'Turno Mañana',
            'Turno Medio Día',
            'Turno Noche',
            'Cumple/No Cumple',]

    for i,titulo in enumerate(titulos):
        celda = table.rows[0].cells[i]
        parrafo = celda.paragraphs[0]
        run = parrafo.add_run(titulo)
        run.bold = True
        run.font.size = Pt(10)

    for (index,conteos),codigos in zip(enumerate(list_counts),list_codigos):
        if index == 0:
            tipicidad = 'Tipico'
        else:
            tipicidad = 'Atipico'
        for i in range(len(codigos)):
            row = table.add_row().cells
            row[0].text = codigos[i]
            row[1].text = tipicidad
            row[2].text = str(50-conteos[i][0])
            row[3].text = str(50-conteos[i][1])
            row[4].text = str(50-conteos[i][2])
            if sum(conteos[i]) == 0:
                row[5].text = 'CUMPLE'
            else:
                row[5].text = 'NO CUMPLE'

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    path_word = os.path.join(ruta_eyb, "Reporte de Embarque y Desembarque.docx")
    doc.save(path_word)