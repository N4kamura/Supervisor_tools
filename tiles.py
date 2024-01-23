from openpyxl import load_workbook
import docx
from docx.shared import Pt,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def tile_report(ruta):
    ruta_tile = os.path.join(ruta,"7.- Informacion de Campo","Longitud de Cola")
    files = ["Atipico","Tipico"]
    listas = [[] for _ in range(2)]

    for index,file in enumerate(files):
        ruta_tipicidad = os.path.join(ruta_tile,file)
        try:
            list_excels = os.listdir(ruta_tipicidad)
        except FileNotFoundError:
            print(f"No hay datos en la carpeta de {file}")
            listas[index] = []
            list_excels = []

        for excel in list_excels:
            path_excel = os.path.join(ruta_tipicidad,excel)

            if not "~$" in excel:
                if file =="Tipico":
                    listas[0].append(path_excel)
                else:
                    listas[1].append(path_excel)

    ####################
    # Counter function #
    ####################

    slices_turnos = (
                slice("D20","K59"), #Turno mañana
                slice("O20","V59"), #Turno medio día
                slice("AK20","AR59"), #Turno noche
            )

    list_conteos = []
    list_codigos = []
    for index,lista in enumerate(listas): #0: Tipico / 1: Atipico
        conteo_excel = []
        codigo_excel = []
        if index == 0:
            print("********** Revisando Longitudes de Cola Tipicos **********")
        else:
            print("********** Revisando Longitudes de Cola Atipicos **********")
        for index2,excel_path in enumerate(lista):
            wb = load_workbook(excel_path,read_only=True,data_only=True)
            ws = wb['Base Data']
        
            conteos = []
            for rebanada in slices_turnos:
                data = list(celda.value for fil in ws[rebanada] for celda in fil)
                contador = 0
                for valor in data:
                    if valor != None:
                        contador+=1
                conteos.append(contador) #(30,30,30)
            conteo_excel.append(conteos)
            _,name_excel = os.path.split(excel_path)
            patron = r"^[0-9]+_([A-Z]+-[0-9]+)"
            coincidencia = re.search(patron,name_excel)
            if coincidencia:
                codigo = coincidencia.group(1)
                codigo = codigo[:2]+'-'+codigo[2:]
            codigo_excel.append(codigo)

        list_conteos.append(conteo_excel)
        list_codigos.append(codigo_excel)

    ########################
    # Creating Word Report #
    ########################

    doc = docx.Document()
    doc.add_heading("REPORTE LONGITUD DE COLA")
    table = doc.add_table(rows=1,cols=6,style="Table Grid")
    section = doc.sections[0]

    for i in range(3):
        for cell in table.columns[i+1].cells:
            cell.width = Inches(1)

    titulos = ['Código',
            'Tipicidad',
            'Turno Mañana',
            'Turno Medio Día',
            'Turno Noche',
            'Cumple/No Cumple']


    for i,titulo in enumerate(titulos):
        celda = table.rows[0].cells[i]
        parrafo = celda.paragraphs[0]
        run = parrafo.add_run(titulo)
        run.bold = True
        run.font.size = Pt(10)
    
    for (index,codigo_excel),conteos in zip(enumerate(list_codigos),list_conteos):
        if index == 0:
            tipicidad = 'Tipico'
        else:
            tipicidad = 'Atipico'
        for i in range(len(codigo_excel)):
            row = table.add_row().cells
            row[0].text = codigo_excel[i]
            row[1].text = tipicidad
            row[2].text = str(conteos[i][0])
            row[3].text = str(conteos[i][1])
            row[4].text = str(conteos[i][2])
            if conteos[i][0]>=30:
                if conteos[i][1]>=30:
                    if conteos[i][2]>=30:
                        row[5].text = 'CUMPLE'
                    else:
                        row[5].text = 'NO CUMPLE'
                else:
                    row[5].text = 'NO CUMPLE'
            else:
                row[5].text = 'NO CUMPLE'

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    path_word = os.path.join(ruta_tile, "Reporte de Longitud de Cola.docx")

    doc.save(path_word)