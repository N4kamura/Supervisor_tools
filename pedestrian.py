from openpyxl import load_workbook
from openpyxl.styles import PatternFill, NamedStyle, Border, Side
import numpy as np
import shutil
import docx
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import os
import re
from reading import read_excel_peatonal
import math

def _put_format(table,row,col): #Just bold
    paragraph = table.cell(row,col).paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True

def _join_cells(table,row1,col1,row2,col2):
    table.rows[row1].cells[col1].merge(table.rows[row2].cells[col2])

def peatonal(entregable_path):
    ruta = entregable_path
    
    listas = [[] for _ in range(4)]

    ruta_N2 = os.path.join(ruta,"7.- Informacion de Campo","Peatonal") #SE CAMBIA
    files = ["Atipico",
            "Atipico (Protransito)",
            "Tipico",
            "Tipico (Protransito)",]
    
    for file in files:
        ruta_penultima = os.path.join(ruta_N2,file)
        hojasexcel = os.listdir(ruta_penultima)
        for hojaexcel in hojasexcel:
            ruta_ultima = os.path.join(ruta_penultima,hojaexcel)

            if not "~$" in hojaexcel:
                        
                if file == "Tipico": #0
                        listas[0].append(ruta_ultima)
                            
                elif file == "Atipico": #1
                        listas[1].append(ruta_ultima)
                        
                elif file == "Tipico (Protransito)": #2
                        listas[2].append(ruta_ultima)
                        
                elif file == "Atipico (Protransito)":
                    listas[3].append(ruta_ultima)

    archivos = [[] for _ in range(2)] #LISTA DE 2
    
    for veh2 in listas[2]: #COMPARACIÓN DE TIPICO CON EL DE PROTRÁNSITO 
        for veh0 in listas[0]:
            nombre_veh2= os.path.basename(veh2)
            nombre_veh0= os.path.basename(veh0)
            if nombre_veh2 != nombre_veh0 and nombre_veh2.replace("_REV","") == nombre_veh0:
                archivos[0].append((veh2,veh0))

    for veh3 in listas[3]: #COMPARACIÓN DE ATIPICO CON EL DE PROTRÁNSITO 
        for veh1 in listas[1]:
            nombre_veh3= os.path.basename(veh3)
            nombre_veh1= os.path.basename(veh1)

            if nombre_veh3 != nombre_veh1 and nombre_veh3.replace("_REV","") == nombre_veh1:
                archivos[1].append((veh3,veh1))

    data_row = [[] for _ in range(2)] #LISTA DE 2 CREA UNA LISTA

    inicio_lectura = time.time()
    
    for j in range(2):
        if j == 0:
            print("#### Reportes Tipicos ####")
        else:
            print("#### Reportes Atipicos ####")

        for index,arch in enumerate(archivos[j]): #el archivo es un par "j" entra a tupla
            if j==0:
                print(f"Comparando Tipicos ({index+1}/{len(archivos[j])})")
            else:
                print(f"Comparando Atipicos ({index+1}/{len(archivos[j])})")
            #0: Supervisor / 1: Consultor
            data_consultor,fecha, giros, _ = read_excel_peatonal(arch[1], False) #pares de archivos - tupla
            data_supervisor,_,giros_compare, quarter_hour = read_excel_peatonal(arch[0], True) #"_" para solo requerir dato de la primera entrada

            #Giros:
            giros_cons = []
            for giro in giros:
                if isinstance(giro, int):
                    giros_cons.append(giro)

            giros_sup = []
            for giro in giros_compare:
                if isinstance(giro, int):
                    giros_sup.append(giro)

            stop = 0
            for (i, giro_cons), giro_sup in zip(enumerate(giros_cons), giros_sup):
                if giro_cons == giro_sup:
                    pass
                else:
                    stop += 1
            check = True
            if stop != 0:
                check = False
                
            ruta_destino = arch[0].replace("_REV.xlsm","_REP.xlsx")
            
            container_file = os.path.dirname(arch[1])
            _,tipicidad = os.path.split(container_file)

            resultados = []

            ##########################
            # LECTURA DE INFORMACIÓN #
            ##########################
            
            A = []
            #LEE LA HOJA Y LUEGO LAS CELDAS CON MASCARAS
            for data_con,data_sup in zip(data_consultor,data_supervisor): #lectura paralela
                mascara = data_sup!=0 #define el verdadero,boleano,verdadero-falso
                resultado = np.zeros_like(data_con)
                resultado[mascara] = data_con[mascara] - data_sup[mascara]
                resultados.append(resultado)

                A.extend(np.abs(resultado[mascara].tolist()))
            
            path_formato = "./images/Formato_Peatonal.xlsx"
            current_file, name_excel = os.path.split(ruta_destino)
            destiny_file = os.path.join(current_file,'Reportes')
            final_route = os.path.join(destiny_file,name_excel)
            if not os.path.exists(destiny_file):
                os.makedirs(destiny_file)
            shutil.copyfile(path_formato,final_route) #copia el contenido del archivo path formato a otro archivo ubicado en la ruta_destino
            wb = load_workbook(final_route) # es la ruta del archivo

            #Definición de formatos:
            porcentaje_style = NamedStyle(name="porcentaje")
            porcentaje_style.border = Border(left=Side(style='thin'),
                                            right=Side(style='thin'),
                                            top=Side(style='thin'),
                                            bottom=Side(style='thin'),)
            ws = wb['Data Peatonal']
            for i in range(len(giros)):
                celda = ws.cell(row=19, column=i+12)
                celda.value = giros[i]

            for result in resultados:
                for k, fila in enumerate(result):
                    for l, valor in enumerate(fila):
                        celda = ws.cell(row=k+20, column=l+12)
                        if valor == 0: celda.value = None
                        else: celda.value = valor

                        celda.style = porcentaje_style
            
            wb.save(final_route)
            _, nombre_archivo = os.path.split(ruta_destino)
            patron1 = r"([A-Z]+[0-9]+)"
            patron2 = r"([A-Z]+-[0-9]+)"
            coincidencia1 = re.search(patron1, nombre_archivo)
            coincidencia2 = re.search(patron2, nombre_archivo)

            if coincidencia1:
                codigo = coincidencia1.group(1)
                codigo = codigo[:2]+'-'+codigo[2:]
            elif coincidencia2:
                codigo = coincidencia2.group(1)
            else:
                print(f"El archivo de excel no posee Código de intersección:\n{nombre_archivo}")
                codigo = "ERROR"
    
            data_row[j].append((codigo, fecha[:-9], tipicidad, A, quarter_hour, check))

    final_lectura = time.time()

    duracion_lectura = final_lectura-inicio_lectura

    print(f"Processing Time - Reading: {duracion_lectura:.2f} seconds.")

    ########################
    # CREACIÓN DEL REPORTE #
    ########################

    inicio_word = time.perf_counter()

    doc = docx.Document()
    doc.add_heading(f"REPORTE FLUJO PEATONAL")
    titulos = ['Clase', 'Frecuencia','Frecuencia Relativa','Frecuencia Acumulada']

    for i, data in enumerate(data_row):
        if len (data) == 0: #No existe data para típico o atípico
            continue

        for dt in data:
            if dt[5] == False:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("Los giros de Protránsito no coinciden con los giros de EDSA.")
                continue
        
            CODIGO      = dt[0]
            FECHA       = dt[1]
            TIPICIDAD   = dt[2]
            LISTA       = dt[3]
            HORA        = dt[4]

            #Tratamiento de la lista completa
            num_datos = len(LISTA)
            num_clases = int(1+3.322*(num_datos**(1/3)))
            rango = max(LISTA) - min(LISTA)
            tamano_clase = rango / num_clases
            tamano_clase_entero = math.ceil(tamano_clase)

            #limite_inferior = min(LISTA)
            limite_inferior = 1
            clases = []
            for _ in range(num_clases):
                limite_superior = limite_inferior + tamano_clase_entero
                clases.append((limite_inferior, limite_superior))
                limite_inferior = limite_superior

            frecuencias = [0]*num_clases
            for dato in LISTA:
                for j, clase in enumerate(clases):
                    if clase[0] <= dato < clase[1]:
                        frecuencias[j] += 1
                        break
            
            if sum(frecuencias) == 0:
                frec_relativa = [0]*len(frecuencias)
            else:
                frec_relativa = [f / sum(frecuencias) for f in frecuencias]
            frec_acumulada = [sum(frec_relativa[:i+1]) for i in range(num_clases)]

            table = doc.add_table(rows=1, cols=4, style='Table Grid')
            table.cell(0,0).text = 'Código'
            table.cell(0,1).text = 'Fecha'
            table.cell(0,2).text = 'Tipicidad'
            table.cell(0,3).text = 'Hora'

            for j in range(4):
                _put_format(table,0,j)

            row = table.add_row().cells
            row[0].text = CODIGO
            row[1].text = str(FECHA)
            row[2].text = str(TIPICIDAD)
            horas_inicio = HORA//4
            minutos_inicio = int(((HORA/4-HORA//4))*60)
            horas_fin = (HORA+1)//4
            minutos_fin = int((((HORA+1)/4-(HORA+1)//4))*60)
            row[3].text = f"{horas_inicio:02}:{minutos_inicio:02} - {horas_fin:02}:{minutos_fin:02}"

            row = table.add_row().cells
            for i, titulo in enumerate(titulos):
                row[i].text = titulo

            for clase, frecuencia, relativa, acumulada, in zip(clases,frecuencias, frec_relativa, frec_acumulada):
                row = table.add_row().cells
                row[0].text = str(f"{int(clase[0])} - {int(clase[1])}")
                row[1].text = str(frecuencia)
                row[2].text = str(int(relativa*100))+'%'
                row[3].text = str(int(acumulada*100))+'%'

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            _put_format(table,2,0)
            _put_format(table,2,1)
            _put_format(table,2,2)
            _put_format(table,2,3)
            
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break()
                        
    #MODIFICAMOS LA UBICACIÓN DEL INFORME 
    try:
        _,nombre_del_archivo = os.path.split(entregable_path)
        direccion_archivo = os.path.join(ruta_N2,f"INFORME_GENERAL_PEATONAL {nombre_del_archivo}.docx") #JUNTO EL NOMBRE DEL DOC CON LA RUTA DE DOCUMENTO , agregar nombre del entregable
        doc.save(direccion_archivo)
        print("Informe general guardado correctamente")    
    except IndexError:
        print(f"No hay datos para generar un informe general.")
    
    final_word = time.time()
    duracion_word = final_word-inicio_word
    print(f"Processing Time - Word: {duracion_word:.2f} segundos.")
    print("#### Comparación de flujogramas peatonales finalizado ####")